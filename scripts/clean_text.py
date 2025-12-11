import re
from docx import Document


def clean_docx_text(file_path: str):
    document = Document(file_path)
    raw_text = []
    for paragraph in document.paragraphs:
        raw_text.append(paragraph.text)

    cleaned = raw_text[0].replace("**ОГРАЖДАЮЩАЯ АКТУАЛЯЦИЯ**", "")
    # Убираем индексы в таблицах (рандомные числа)
    cleaned = re.sub(r"\|\s*\d+\s*(?:\|\s*){2,}.*", "", cleaned)
    # Разделители таблиц
    cleaned = re.sub(r"\|[\-]+\|", " ", cleaned)
    # MD
    for char in ["*", "#", "|"]:
        cleaned = cleaned.replace(char, " ")
    # Нормализация реальных \n
    cleaned = re.sub(r"\n\s*\n", "\n", cleaned)
    # \n в тексте
    cleaned = re.sub(r"\\n", "", cleaned)
    # Лишние пробелы
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    # Пустые таблицы
    cleaned = re.sub(r"\n\s*\d+\s*\n", "\n", cleaned)

    # with open(f"{file_path}.txt", 'w', encoding='utf-8') as f:
    #     f.write(cleaned.strip())

    return cleaned.strip()


if __name__ == "__main__":
    clean_text = clean_docx_text("../files/ADC_8.docx")
    raw_text = Document("../files/ADC_8.docx").paragraphs[0].text

    print(f"Длина raw текса: {len(raw_text)}")
    print(f"Длина чистого текса: {len(clean_text)}")
    print(clean_text[-100:])
    print(raw_text[-100:])
